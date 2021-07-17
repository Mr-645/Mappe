import sys
import os
from PyQt5.QtWidgets import *
from PyQt5.QtGui import *
from PyQt5.QtCore import *
import json
import docx
from PyQt5.uic import loadUi, loadUiType
import sqlite3
import itertools
import zipfile
import re
from datetime import datetime, date
from dateutil.relativedelta import relativedelta
from glob import glob
from docx2python import docx2python

# from app2 import Ui_MainWindow

From_Main, _ = loadUiType(os.path.join(os.path.dirname(__file__), "C:/Users/drift/OneDrive/PROJECTS/Mappe/testing_ground/app2.ui"))

class MainWindow(QMainWindow, From_Main):


# class Window(QMainWindow, Ui_MainWindow):

    def __init__(self, parent=None):

        # super(MainWindow, self).__init__()
        super().__init__(parent)

        self.show()

        self.setupUi(self)

        # ====== Functionality code starts here ======
        self.setWindowIcon(QIcon('C:/Users/drift/OneDrive/PROJECTS/Mappe/icons/Mappe_icon.png'))
        self.statusBar.showMessage(">")

        self.actionCreate_TQR_and_PC.triggered.connect(self.create_TQR_and_PC)
        self.actionCreate_new_SOP.triggered.connect(self.create_new_sop)
        self.actionRescan_for_main_table.triggered.connect(self.rescan_dir_for_DB)
        self.actionRefresh_main_table.triggered.connect(self.table_refresh)

        self.connectButtons()

        self.load_global_variables()

        self.table_refresh()

        self.make_header()

        self.header.setSectionsClickable(True)
        self.tableWidget.setSortingEnabled(True)
        self.header.setSortIndicatorShown(True)

        refreshAction = QAction(QIcon("C:/Users/drift/OneDrive/PROJECTS/Mappe/icons/refresh.png"), "Refresh table", self)
        refreshAction.triggered.connect(self.table_refresh)
        self.tableWidget.addAction(refreshAction)

        self.tableWidget.viewport().installEventFilter(self)


    def eventFilter(self, source, event):
        if event.type() == QEvent.MouseButtonRelease:
            row = self.tableWidget.currentRow()
            col = self.tableWidget.currentColumn()
            print(row,col)
        return QObject.event(source, event)

    def connectButtons(self):
        self.actionSystem_parameters.triggered.connect(self.open_pth_edit_dialogue)
        self.actionExit.triggered.connect(self.close)
        self.tableWidget.itemSelectionChanged.connect(self.handle_selection)

        self.combo = QComboBox()
        self.toolBar.addWidget(self.combo)
        self.combo.addItems(["Open selected SOP file(s)", "Open Task-Qualification-Report(s) of selected file(s)", "Open Procedure-Checklist(s) of selected file(s)"])

        self.go_button = QPushButton("Go")
        self.toolBar.addWidget(self.go_button)
        self.go_button.setIcon(QIcon("C:/Users/drift/OneDrive/PROJECTS/Mappe/qt-designer-python/sample_editor/ui/resources/file-open.png"))
        self.go_button.clicked.connect(self.open_sop_link)


    def load_global_variables(self):
        path = "C:/Users/drift/OneDrive/PROJECTS/Mappe/testing_ground/global_var_2.json"
        with open(path, 'r') as config_file_read:
            self.loaded_vars = json.load(config_file_read)

        self.sop_path = self.loaded_vars["sop_doc_folder"]
        self.tqr_path = self.loaded_vars["task_qual_rep_folder"]
        self.pc_path = self.loaded_vars["proc_checklist_folder"]
        self.temp_path = self.loaded_vars["template_doc_folder"]
        self.main_table_path = self.loaded_vars["main_table_data_file"]
        self.file_expiry = self.loaded_vars["file_expiry_time_months"]


    def fill_table_sqlite(self, filter_params):
        self.tableWidget.clear()
        self.tableWidget.setRowCount(0)
        self.statusBar.showMessage(">")
        
        try:
            connection = sqlite3.connect(self.main_table_path)
        except:
            self.statusBar.showMessage("> ERROR: Can't open main-table database.")
            return None

        query2 = "PRAGMA table_info(main_table_db)"
        result2 = connection.execute(query2)

        header_list = result2.fetchall()
        self.header_list_2 = []

        for i in header_list:
            self.header_list_2.append(i[1])

        if filter_params == None or filter_params == False:
            query1 = "SELECT * FROM main_table_db"
        else:
            query_suffix = ""
            for (p1, p2), i in itertools.zip_longest(filter_params, range(len(list(filter_params)))):
                if p2 == "" or p2 == "'":
                    p2 = "**"
                query_suffix += f"\"{str(self.header_list_2[p1])}\" glob \"{p2}\""
                if ( i != len(filter_params)-1 ):
                    query_suffix += f" and "
            
            query1 = f"SELECT * FROM main_table_db WHERE {query_suffix}"

        try:
            result1 = connection.execute(query1)
        except:
            self.statusBar.showMessage("> There's a syntax error in the Filter header")
            result1 = connection.execute("SELECT * FROM main_table_db")

        self.tableWidget.setColumnCount(len(self.header_list_2))
        self.tableWidget.setHorizontalHeaderLabels(list(self.header_list_2))

        for row_number, row_data in enumerate(result1):
            self.tableWidget.insertRow(row_number)
            for column_number, data in enumerate(row_data):
                if (self.header_list_2[column_number] == "SOP expiry status") and (data == "expired"):
                    self.tableWidget.setItem(row_number, column_number, QTableWidgetItem(str(data)))
                    # QTableWidgetItem
                    self.tableWidget.item(row_number, column_number).setTextAlignment(Qt.AlignCenter)
                    self.tableWidget.item(row_number, column_number).setBackground(QColor(255,0,0))
                    self.tableWidget.item(row_number, column_number).setForeground(Qt.yellow)
                else:
                    self.tableWidget.setItem(row_number, column_number, QTableWidgetItem(str(data)))
                    # self.tableWidget.item(row_number, column_number).setBackground(QColor(255,255,255))

        connection.close()

        self.tableWidget.resizeColumnsToContents()
        self.tableWidget.resizeRowsToContents()

    def handle_selection(self):
        indexes = self.tableWidget.selectionModel().selectedRows()
        self.statusBar.showMessage(f"You have selected {len(indexes)} document(s)")
        self.selection_list = []
        for index in indexes:
            self.selection_list.append([
                self.tableWidget.model().data(self.tableWidget.model().index(index.row(), 0)), # ID
                self.tableWidget.model().data(self.tableWidget.model().index(index.row(), 1)), # SOP name
                self.tableWidget.model().data(self.tableWidget.model().index(index.row(), 2)), # last mod date
                self.tableWidget.model().data(self.tableWidget.model().index(index.row(), 3)), # rev number
                self.tableWidget.model().data(self.tableWidget.model().index(index.row(), 4)), # issue date
                self.tableWidget.model().data(self.tableWidget.model().index(index.row(), 5)), # review date
                self.tableWidget.model().data(self.tableWidget.model().index(index.row(), 6)), # expiry status
                self.tableWidget.model().data(self.tableWidget.model().index(index.row(), 7)), # SOP file link
                self.tableWidget.model().data(self.tableWidget.model().index(index.row(), 8)), # TQR file link
                self.tableWidget.model().data(self.tableWidget.model().index(index.row(), 9)) # PC file link
            ])
    
    def table_refresh(self):
        self.tableWidget.clear()
        self.tableWidget.setRowCount(0)
        self.tableWidget.setColumnCount(0)
        self.fill_table_sqlite(None)
    
    def create_pop_up(self, type, maintext, informtext, wintitle, detailtext):
        msg = QMessageBox()
        msg.setIcon(type)
        msg.setText(maintext)
        msg.setInformativeText(informtext)
        msg.setWindowTitle(wintitle)
        msg.setDetailedText(detailtext)
        msg.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)
        msg.setWindowIcon(QIcon('C:/Users/drift/OneDrive/PROJECTS/Mappe/icons/Mappe_icon.png'))

        returnValue = msg.exec()
        return returnValue

    def create_TQR_and_PC(self):
        self.statusBar.showMessage("> Progress: ")

        self.progressBar_instatusBar_2 = QProgressBar()
        self.statusBar.addPermanentWidget(self.progressBar_instatusBar_2)
        self.progressBar_instatusBar_2.setGeometry(0, 0, 00, 00)
        # self.progressBar_instatusBar_2.setGeometry(400, 5, 50, 10)
        self.progressBar_instatusBar_2.setValue(0)
        
        self.progressBar_instatusBar_1 = QProgressBar()
        self.statusBar.addPermanentWidget(self.progressBar_instatusBar_1)
        self.progressBar_instatusBar_1.setGeometry(0, 0, 00, 00)
        self.progressBar_instatusBar_1.setGeometry(70, 5, 300, 20)
        self.progressBar_instatusBar_1.setValue(0)

        self.progressBar_instatusBar_2.setGeometry(400, 10, 100, 10)
        # self.tableWidget.horizontalHeader().sectionClicked.connect(self.on_header_clicked)
        
        try:
            self.selection_list
        except:    
            self.selection_list = []
        
        sorted_selection_list = ""
        for x in self.selection_list:
            sorted_selection_list += f"{x[1]},\n"

        # returnValue = msg.exec()
        type = QMessageBox.Question
        maintext = "Do you want to create \t\t\t\t\t\t\t \nTask Qualification Report(s) and \nProcedure checklists \nof the selected SOPs?"
        informtext = f"You have selected {len(self.selection_list)} SOP(s)"
        wintitle = "Create Task Qualification Report(s)"
        detailtext = f"---These are the selected SOPs:\n{sorted_selection_list}"

        returnValue = self.create_pop_up(type, maintext, informtext, wintitle, detailtext)

        # if returnValue == QMessageBox.Ok:
        if returnValue == 1024:
            for number, i in enumerate(self.selection_list):
                self.fill_PC_and_TQR(i[7], i[1], i[3], i[0]) #pass SOP name & path to the function below
                self.progressBar_instatusBar_1.setValue((( (number+1) * 100 ) // len(self.selection_list)))

        self.progressBar_instatusBar_2.hide()
        self.progressBar_instatusBar_1.hide()

        del self.progressBar_instatusBar_2
        del self.progressBar_instatusBar_1
        
        self.table_refresh()
        self.statusBar.showMessage("> Done ")
    
    def fill_PC_and_TQR(self, correlating_sop, sop_name, rev_num, sop_code):
        doc_result = docx2python(correlating_sop)
        self.progressBar_instatusBar_2.setValue(5)

        out_list_yellow = []
        out_dict = {}

        for i in range(len(doc_result.body)):
            for j in range(len(doc_result.body[i])):
                for k in range(len(doc_result.body[i][j])):
                    out_list_yellow.append(doc_result.body[i][j][k])
        self.progressBar_instatusBar_2.setValue(10)

        list2 = []
        # This entire for-loop is for getting the starting point
        for i in range(0, len(out_list_yellow), 1):
            try:
                if (out_list_yellow[i][0] == ('Machine Safety Awareness')):
                    start_point = i
            except:
                None
        self.progressBar_instatusBar_2.setValue(15)

        # Get rid of the empty (or space character containing), and photo target list entries
        for i in range(start_point, len(out_list_yellow), 1):
            if list(filter(None, out_list_yellow[i])): #removes empties
                if ("----media/image" not in out_list_yellow[i][0]): #removes "----medi.."
                    list2.append(out_list_yellow[i])
        self.progressBar_instatusBar_2.setValue(20)

        # Get rid of the tab characters in all the list entries in every nested list
        for i in range(len(list2)):
            for j in range(len(list2[i])):
                list2[i] = ([s.replace('\t', '') for s in list2[i]])
        self.progressBar_instatusBar_2.setValue(25)

        # Put the data into a dictionary 
        # -Major steps
        for i in range(0, len(list2), 3):
            out_dict[list2[i][0]] = None
        self.progressBar_instatusBar_2.setValue(30)

        problematic_sops = []
        # -combined key points and reasons
        for i, j in itertools.zip_longest(range(1, len(list2), 3), range(len(list(out_dict.keys())))):
            if i == None or j == None:
                problematic_sops.append(sop_name)
            else:
                try:
                    out_dict.update({ list(out_dict.keys())[j] : [list2[i],list2[i+1]] })
                except:
                    continue
        self.progressBar_instatusBar_2.setValue(35)

        type = QMessageBox.Critical
        maintext = "Some of the SOP documents you've selected, have problems."
        informtext = "REMEMBER: if any cell in the SOP (except\n in the 'Photos/Diagrams' column) is empty, I'll throw an ERROR!"
        wintitle = "Problematic SOPs"
        detailtext = "\n".join(problematic_sops)
            
        if len(problematic_sops) != 0:
            self.create_pop_up(type, maintext, informtext, wintitle, detailtext)
        self.progressBar_instatusBar_2.setValue(40)

        out_dict_keys_list = list(out_dict.keys())
        out_dict_values_list = list(out_dict.values())
        self.progressBar_instatusBar_2.setValue(45)

        #remove '&' and replace it with '&amp;' to prevent an error with LXML
        out_dict_keys_list = ([s.replace('&', '&amp;') for s in out_dict_keys_list])

        for i in range(len(out_dict_values_list)):
            for j in range(len(out_dict_values_list[i])):
                out_dict_values_list[i][j] = ([s.replace('&', '&amp;') for s in out_dict_values_list[i][j]])
        self.progressBar_instatusBar_2.setValue(50)

        out_dict_keys_list = list(filter(None, out_dict_keys_list)) #removes empty entries

        # Create the procedure checklist from the SOP and the template file
        pc_template = self.temp_path + "/Procedure_checklist_template.docx"
        created_PC_file = f"{self.pc_path}/PC_{sop_name}.docx"

        zin = zipfile.ZipFile (pc_template, 'r')
        self.progressBar_instatusBar_2.setValue(55)

        try:
            zout = zipfile.ZipFile (created_PC_file, 'w')
        except:
            self.create_pop_up(QMessageBox.Critical, "Error: I can't do anything with the file for some reason...", str(sys.exc_info()[0]), "Error", None)
            return None

        self.progressBar_instatusBar_2.setValue(60)
        # zout = zipfile.ZipFile (created_PC_file, 'w')

        for item in zin.infolist():
            buffer = zin.read(item.filename)
            if (item.filename == 'word/document.xml'):
                res = buffer.decode("utf-8")
                res = res.replace("thesopname", sop_name)
                res = res.replace("tasktitle", out_dict_keys_list[3])
                for i in range(4, len(out_dict_keys_list)):                    
                    if (out_dict_keys_list[i])[1] == ")": #removes the list letter e.g. b) from the text
                        out_dict_keys_list[i] = (out_dict_keys_list[i])[2:]
                    res = res.replace(f"ms{i-3}_", out_dict_keys_list[i])
                buffer = res.encode("utf-8")
            elif (item.filename == 'word/header1.xml'):
                res = buffer.decode("utf-8")
                res = res.replace("thesopname", sop_name)
                buffer = res.encode("utf-8")
            elif (item.filename == 'word/footer1.xml'):
                res = buffer.decode("utf-8")
                res = res.replace("thesopname", sop_name)
                res = res.replace("todaysdate", date.today().strftime("%d %B %Y"))
                buffer = res.encode("utf-8")
            zout.writestr(item, buffer)
        zout.close()
        zin.close()
        self.progressBar_instatusBar_2.setValue(65)

        #Remove the unedited columns
        doc = docx.Document(created_PC_file)
        table = doc.tables[0]
        self.progressBar_instatusBar_2.setValue(70)

        for n in range(80 - len(out_dict_keys_list)):
            row = table.rows[-1] #start from the end of the list
            # remove_row(table, row)
            tbl = table._tbl
            tr = row._tr
            tbl.remove(tr)

        doc.save(created_PC_file)
        self.progressBar_instatusBar_2.setValue(75)

        # Create the task qualification report from the SOP and the template file
        tqr_template = self.temp_path + "/Task_Qualification_report_template.docx"
        created_tqr_file = f"{self.tqr_path}/TQR_{sop_name}.docx"

        self.progressBar_instatusBar_2.setValue(80)

        zin = zipfile.ZipFile (tqr_template, 'r')
        zout = zipfile.ZipFile (created_tqr_file, 'w')
        for item in zin.infolist():
            buffer = zin.read(item.filename)
            if (item.filename == 'word/document.xml'):
                res = buffer.decode("utf-8")
                res = res.replace("thesopname", sop_name)
                for i in range(0, len(out_dict_keys_list)):
                    if out_dict_values_list[i] == None:
                        pass
                    else:
                        if (out_dict_keys_list[i])[1] == ")": #removes the list letter e.g. b) from the text
                            out_dict_keys_list[i] = (out_dict_keys_list[i])[2:]

                        res = res.replace(f"timms{i+1}e", out_dict_keys_list[i])
                        res = res.replace(f"timkp{i+1}e", self.process_tqr_kp_r_data(out_dict_values_list[i],0))
                        res = res.replace(f"timr{i+1}e", self.process_tqr_kp_r_data(out_dict_values_list[i],1))
                buffer = res.encode("utf-8")
            elif (item.filename == 'word/header1.xml'):
                res = buffer.decode("utf-8")
                res = res.replace("thesopcode", sop_code)
                res = res.replace("thesoprevnum", f"Revision number: {rev_num}")
                res = res.replace("thesopname", sop_name)
                buffer = res.encode("utf-8")
            elif (item.filename == 'word/header2.xml'):
                res = buffer.decode("utf-8")
                res = res.replace("thesopcode", sop_code)
                res = res.replace("thesoprevnum", f"Revision number: {rev_num}")
                res = res.replace("thesopname", sop_name)
                buffer = res.encode("utf-8")
            zout.writestr(item, buffer)
        zout.close()
        zin.close()
        self.progressBar_instatusBar_2.setValue(85)

        #Remove the unedited columns
        doc = docx.Document(created_tqr_file)
        table = doc.tables[1]
        self.progressBar_instatusBar_2.setValue(90)

        for n in range(88 - len(out_dict_keys_list)):
            row = table.rows[-1] #start from the end of the list
            # remove_row(table, row)
            tbl = table._tbl
            tr = row._tr
            tbl.remove(tr)

        doc.save(created_tqr_file)

        try:
            connection = sqlite3.connect(self.main_table_path)
            self.progressBar_instatusBar_2.setValue(95)
        except:
            self.statusBar.showMessage("> ERROR: Can't open main-table database.")
            return None

        exe_str = f"UPDATE main_table_db SET \"Task qualification report file link\" = \"{created_tqr_file}\", \"Procedure checklist file link\" = \"{created_PC_file}\" WHERE \"SOP name\" = \"{sop_name}\""
        connection.execute(exe_str)
        connection.commit()
        connection.close()
        self.progressBar_instatusBar_2.setValue(100)

    def process_tqr_kp_r_data(self, input_list, col): #col=0 (key point) or col=1 (reasons)
        indent = "                                  "
        output = ""

        for i in range(len(input_list[col])):
            try:
                if (input_list[col][i])[1] == ")": #removes the list letter e.g. b) from the text
                    input_list[col][i] = (input_list[col][i])[2:]
            except:
                continue

            if len(input_list[col]) > 1:
                if i == 0:
                    output += f"- {input_list[col][i]}<w:br/>"
                else: output += f"{indent}- {input_list[col][i]}<w:br/>"
            else:
                output += f"{input_list[col][i]}"

        return output

    def create_new_sop(self):
        dialogue1 = NewSOPDialogue(self)
        dialogue1.exec()

    def rescan_dir_for_DB(self):
        dialogue1 = DirRescanDialogue(self)
        dialogue1.exec()

    def open_pth_edit_dialogue(self):
        dialogue1 = FileFolderPathDialogue(self)
        dialogue1.exec()
    
    def flash(self):
        # print(f"flash: {interval}")
        n = self.statusBar.palette()
        n.setColor(self.statusBar.backgroundRole(), Qt.red)
        n.setColor(self.statusBar.foregroundRole(), Qt.white)
        
        o = self.statusBar.palette()
        o.setColor(self.statusBar.backgroundRole(), Qt.cyan)
        o.setColor(self.statusBar.foregroundRole(), Qt.black)
        
        e = self.statusBar.palette()
        e.setColor(self.statusBar.backgroundRole(), Qt.white)
        e.setColor(self.statusBar.foregroundRole(), Qt.black)

        if self.flash_func_counter % 2 == 0:
            self.statusBar.setPalette(n)
        else:
            self.statusBar.setPalette(o)
        
        self.flash_func_counter += 1

        if self.flash_func_counter == 40:
            self.timer.stop()
            self.statusBar.setPalette(e)
    
    def flash_statbar_message(self, mymessage):

        self.statusBar.setAutoFillBackground(True)
        self.statusBar.showMessage(f"> {mymessage}")
        
        self.flash_func_counter = 0
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.flash)
        self.timer.start(100)

    def open_sop_link(self):
        if (self.combo.currentIndex() == 0):
            for i in self.selection_list:
                if len(i[7]) != 0:
                    os.startfile(i[7])
                else:
                    self.flash_statbar_message("There's nothing there. No file.")

        elif (self.combo.currentIndex() == 1):
            for i in self.selection_list:
                if len(i[8]) != 0:
                    os.startfile(i[8])
                else:
                    self.flash_statbar_message("There's nothing there. No file.")

        elif (self.combo.currentIndex() == 2):
            for i in self.selection_list:
                if len(i[9]) != 0:
                    os.startfile(i[9])
                else:
                    self.flash_statbar_message("There's nothing there. No file.")

    def make_header(self):
        self.header = FilterHeader(self.tableWidget)
        self.tableWidget.setHorizontalHeader(self.header)
        self.header.setFilterBoxes(self.tableWidget.columnCount())
        self.header.filterActivated.connect(self.handleFilterActivated)
        
        # self.header.setSectionsClickable(True)

        self.tableWidget.resizeColumnsToContents()
        self.tableWidget.resizeRowsToContents()

    def handleFilterActivated(self):
        header = self.tableWidget.horizontalHeader()

        filter_list = []
        for index in range(header.count()):
            filter_list.append([index, header.filterText(index)])

        self.tableWidget.setSortingEnabled(False)

        self.fill_table_sqlite(filter_list)

        self.tableWidget.setSortingEnabled(True)

        self.tableWidget.resizeColumnsToContents()
        self.tableWidget.resizeRowsToContents()


class FilterHeader(QHeaderView):
    filterActivated = pyqtSignal()

    def __init__(self, parent=None):
        super().__init__(Qt.Horizontal, parent)
        self._editors = []
        self._padding = 4
        self.setStretchLastSection(False)
        # self.setResizeMode(QHeaderView.Stretch)
        self.setDefaultAlignment(Qt.AlignLeft | Qt.AlignVCenter)
        self.setSortIndicatorShown(False)
        self.sectionResized.connect(self.adjustPositions)
        parent.horizontalScrollBar().valueChanged.connect(self.adjustPositions)

        # self.setSectionsClickable(True) # this allows you to click the header and then go onto sort the table

    def setFilterBoxes(self, count):
        while self._editors:
            editor = self._editors.pop()
            editor.deleteLater()
        for index in range(count):
            # if index == 33:
                # editor = QComboBox(self.parent())
                # editor.returnPressed.connect(self.filterActivated.emit)
                # editor.addItems(["One", "Two"])
            # else:
            editor = QLineEdit(self.parent())
            editor.setPlaceholderText('Filter')
            editor.returnPressed.connect(self.filterActivated.emit)
            self._editors.append(editor)
        self.adjustPositions()

    def sizeHint(self):
        size = super().sizeHint()
        if self._editors:
            height = self._editors[0].sizeHint().height()
            size.setHeight(size.height() + height + self._padding)
        return size

    def updateGeometries(self):
        if self._editors:
            height = self._editors[0].sizeHint().height()
            self.setViewportMargins(0, 0, 0, height + self._padding)
        else:
            self.setViewportMargins(0, 0, 0, 0)
        super().updateGeometries()
        self.adjustPositions()

    def adjustPositions(self):
        if (self.parent().rowCount()) == 0:
            my_offset = 0
        elif (self.parent().rowCount()) in range (1, 10):
            my_offset = 16
        elif (self.parent().rowCount()) in range (10, 100):
            my_offset = 23
        elif (self.parent().rowCount()) in range (100, 1000):
            my_offset = 30
        elif (self.parent().rowCount()) in range (1000, 10000):
            my_offset = 37
        else:
            my_offset = 44
        
        for index, editor in enumerate(self._editors):
            height = editor.sizeHint().height()
            editor.move(self.sectionPosition(index) - self.offset() + my_offset, height + (self._padding // 2))
            # editor.move(self.sectionPosition(index) - self.offset() + 16, height + (self._padding // 2))
            editor.resize(self.sectionSize(index) , height)

    def filterText(self, index):
        if 0 <= index < len(self._editors):
            return self._editors[index].text()
        return ''

    def setFilterText(self, index, text):
        if 0 <= index < len(self._editors):
            self._editors[index].setText(text)

    def clearFilters(self):
        for editor in self._editors:
            editor.clear()


class FileFolderPathDialogue(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        loadUi("C:/Users/drift/OneDrive/PROJECTS/Mappe/testing_ground/doc_dir.ui", self)
        self.setWindowIcon(QIcon('C:/Users/drift/OneDrive/PROJECTS/Mappe/icons/Mappe_icon.png'))
        
        # Initialise OK-Apply-Cancel button functionality
        self.pushButton.clicked.connect(self.save_close_changes)
        self.pushButton_2.clicked.connect(self.save_changes)
        self.pushButton_3.clicked.connect(self.close)

        # Initialise toolbutton functionality
        self.toolButton.clicked.connect(self.choose_path_sops)
        self.toolButton_2.clicked.connect(self.choose_path_task_qual_report)
        self.toolButton_3.clicked.connect(self.choose_path_procedure_checklist)
        self.toolButton_4.clicked.connect(self.choose_path_template_folder)
        self.toolButton_5.clicked.connect(self.choose_path_main_table)

        # Populate field values/data
        self.lineEdit.setText(mywindow.sop_path)
        self.lineEdit_2.setText(mywindow.tqr_path)
        self.lineEdit_3.setText(mywindow.pc_path)
        self.lineEdit_4.setText(mywindow.temp_path)
        self.lineEdit_5.setText(mywindow.main_table_path)
        self.spinBox.setValue(mywindow.file_expiry)


        self.label_1.setText(f"<a href={mywindow.sop_path}>{self.label_1.text()}</a>")
        self.label_2.setText(f"<a href={mywindow.tqr_path}>{self.label_2.text()}</a>")
        self.label_3.setText(f"<a href={mywindow.pc_path}>{self.label_3.text()}</a>")
        self.label_4.setText(f"<a href={mywindow.temp_path}>{self.label_4.text()}</a>")

    def save_changes(self):

        global_variable_dict = {
            "sop_doc_folder": self.lineEdit.text().rstrip("/"),
   	        "task_qual_rep_folder": self.lineEdit_2.text().rstrip("/"),
   	        "proc_checklist_folder": self.lineEdit_3.text().rstrip("/"),
   	        "template_doc_folder": self.lineEdit_4.text().rstrip("/"),
   	        "main_table_data_file": self.lineEdit_5.text().rstrip("/"),
   	        "file_expiry_time_months": self.spinBox.value()
        }

        test_globvar = "C:/Users/drift/OneDrive/PROJECTS/Mappe/testing_ground/global_var_2.json"
        config_file_write = open(test_globvar, 'w')
        json.dump(global_variable_dict, config_file_write)

        mywindow.sop_path = self.lineEdit.text().rstrip("/")
        mywindow.tqr_path = self.lineEdit_2.text().rstrip("/")
        mywindow.pc_path = self.lineEdit_3.text().rstrip("/")
        mywindow.temp_path = self.lineEdit_4.text().rstrip("/")
        mywindow.main_table_path = self.lineEdit_5.text().rstrip("/")
        mywindow.file_expiry = self.spinBox.value()

        mywindow.statusBar.showMessage("Apply: Paths sucessfully updated")

    def save_close_changes(self):
        self.save_changes()
        mywindow.statusBar.showMessage("OK: Paths sucessfully updated")
        self.close()

    def choose_path_sops(self):
        path = QFileDialog.getExistingDirectory(self, 'Select folder containing SOP documents', os.getenv('HOME'))
        if path != "":
            self.lineEdit.setText(path)

    def choose_path_task_qual_report(self):
        path = QFileDialog.getExistingDirectory(self, 'Select folder containing task qualification report documents', os.getenv('HOME'))
        if path != "":
            self.lineEdit_2.setText(path)

    def choose_path_procedure_checklist(self):
        path = QFileDialog.getExistingDirectory(self, 'Select folder containing procedure checklist documents', os.getenv('HOME'))
        if path != "":
            self.lineEdit_3.setText(path)

    def choose_path_template_folder(self):
        path = QFileDialog.getExistingDirectory(self, 'Select folder containing template documents', os.getenv('HOME'))
        if path != "":
            self.lineEdit_4.setText(path)

    def choose_path_main_table(self):
        path = QFileDialog.getOpenFileName(self, 'Select file for populating main table', os.getenv('HOME'), 'SQLite Database(*.db)')[0]
        if path != "":
            self.lineEdit_5.setText(path)


class NewSOPDialogue(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        loadUi("C:/Users/drift/OneDrive/PROJECTS/Mappe/testing_ground/new_sop.ui", self)
        self.setWindowIcon(QIcon('C:/Users/drift/OneDrive/PROJECTS/Mappe/icons/Mappe_icon.png'))

        # Initialise OK-Cancel button functionality

        # self.buttonBox.button(QDialogButtonBox.Ok).clicked.connect(lambda: self.save_changes("ok"))
        # self.buttonBox.button(QDialogButtonBox.Cancel).clicked.connect(lambda: self.save_changes("cancel"))

        self.buttonBox.button(QDialogButtonBox.Ok).clicked.connect(self.save_changes)

    def save_changes(self):
        sop_template = mywindow.temp_path + "/SOP_template.docx"
        created_sop_file = f"{mywindow.sop_path}/{self.lineEdit.text()}.docx"

        zin = zipfile.ZipFile (sop_template, 'r')
        zout = zipfile.ZipFile (created_sop_file, 'w')
        for item in zin.infolist():
            buffer = zin.read(item.filename)
            if (item.filename == 'word/document.xml'):
                res = buffer.decode("utf-8")
                res = res.replace("Tbp1", str(re.sub('&', '&amp;', self.lineEdit.text()))) #lineEdit text to go here
                res = res.replace("The_purpose", str(re.sub('&', '&amp;', self.lineEdit_3.text()))) #lineEdit text to go here
                res = res.replace("The_scope", str(re.sub('&', '&amp;', self.lineEdit_4.text()))) #lineEdit text to go here
                res = res.replace("The_product", str(re.sub('&', '&amp;', self.lineEdit_5.text()))) #lineEdit text to go here
                res = res.replace("The_pre", str(re.sub('&', '&amp;', self.lineEdit_6.text()))) #lineEdit text to go here
                res = res.replace("The_equipment", str(re.sub('&', '&amp;', self.lineEdit_7.text()))) #lineEdit text to go here
                buffer = res.encode("utf-8")
            elif (item.filename == 'word/header2.xml'):
                res = buffer.decode("utf-8")
                res = res.replace("Ttp1", str(re.sub('&', '&amp;', self.lineEdit.text()))) #lineEdit text to go here
                buffer = res.encode("utf-8")
            elif (item.filename == 'word/header4.xml'):
                res = buffer.decode("utf-8")
                res = res.replace("Ttp2", str(re.sub('&', '&amp;', self.lineEdit.text()))) #lineEdit text to go here
                buffer = res.encode("utf-8")
            elif (item.filename == 'word/footer2.xml'):
                res = buffer.decode("utf-8")
                res = res.replace("SOP Code", str(re.sub('&', '&amp;', self.lineEdit_2.text()))) #lineEdit text to go here
                buffer = res.encode("utf-8")
            elif (item.filename == 'word/footer4.xml'):
                res = buffer.decode("utf-8")
                res = res.replace("SOP Code", str(re.sub('&', '&amp;', self.lineEdit_2.text()))) #lineEdit text to go here
                buffer = res.encode("utf-8")
            zout.writestr(item, buffer)
        zout.close()
        zin.close()

        mywindow.statusBar.showMessage(f"New SOP: {str(re.sub('&', '&amp;', self.lineEdit.text()))}, has been created")


class DirRescanDialogue(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        loadUi("C:/Users/drift/OneDrive/PROJECTS/Mappe/testing_ground/dir_scan.ui", self)
        self.setWindowIcon(QIcon('C:/Users/drift/OneDrive/PROJECTS/Mappe/icons/Mappe_icon.png'))

        self.tableWidget_2.setColumnWidth(0, 220)
        self.tableWidget_2.setColumnWidth(1, 130)

        # Initialise OK-Cancel button functionality
        self.pushButton.clicked.connect(self.start_scan_for_changes)
        self.pushButton_2.clicked.connect(self.close)
    
    #get all the necessary info from all the 'need to add' SOPs
    def get_SOP_info_from_doc(self, file):
        doc_result = docx2python(file)

        header_list = []

        for i in range(len(doc_result.header)):
            for j in range(len(doc_result.header[i])):
                for k in range(len(doc_result.header[i][j])):
                    if (doc_result.header[i][j][k][0]) != "":
                        header_list.append(doc_result.header[i][j][k][0])

        header_list2 = []

        for i in range(len(header_list)):
                if 'Authorised' in header_list[i]:
                    header_list2.append(header_list[i])
                elif 'Revision' in header_list[i]:
                    header_list2.append(header_list[i])
                elif 'Issue' in header_list[i]:
                    header_list2.append(header_list[i])
                elif 'Review' in header_list[i]:
                    header_list2.append(header_list[i])

        rev_num = header_list2[1].strip('Revision: ')
        issue_date = header_list2[2].strip('Issue Date: ')
        review_date = header_list2[3].strip('Review Date: ')

        footer_list = []

        for i in range(len(doc_result.footer)):
            for j in range(len(doc_result.footer[i])):
                for k in range(len(doc_result.footer[i][j])):
                    if (doc_result.footer[i][j][k][0]) != "":
                        footer_list.append(doc_result.footer[i][j][k][0])

        sop_code = footer_list[0]
        sop_code = sop_code.split("\tPage" , -1)
        sop_code = [x.replace('\t', '') for x in sop_code]
        sop_code = str(sop_code[0])

        return rev_num, issue_date, review_date, sop_code

    def convert_date(self, timestamp):
        d = datetime.utcfromtimestamp(timestamp)
        return d.strftime('%d/%m/%Y') # e.g. 22/04/2021
    
    def start_scan_for_changes(self):
        mywindow.statusBar.showMessage(f"Started scan for changes")

        #create new list, look through SOP directory, get filenames, put names and URIs into list
        files = []
        for (dir,a,b) in os.walk(mywindow.sop_path):
            files.extend(glob(os.path.join(dir,"*.docx")))
            self.tableWidget_2.setItem(0, 0, QTableWidgetItem(f"Found {len(files)} file(s)"))

        files = ([s.replace('\\', '/') for s in files])

        file_listing = []
        # file_listing_big = []
        for index, i in enumerate(files):
            file_listing.append([(os.path.basename(i)).replace('.docx', ''), self.convert_date(os.path.getmtime(i)), i])
            self.tableWidget_2.setItem(1, 0, QTableWidgetItem(f"Processed {index+1} out of {len(files)} file(s)"))
            self.progressBar_7.setValue((index + 1) // (len(files))*100)

        #compare list(name) contents with document name column in sqlitedb table, create list of 'new' and 'not_present'
        
        try:
            connection = sqlite3.connect(mywindow.main_table_path)
        except:
            self.statusBar.showMessage("> ERROR: Can't open main-table database.")
            return None

        get_names_query = "SELECT \"SOP name\" FROM main_table_db"
        file_name_list = [item[0] for item in connection.execute(get_names_query).fetchall()]
        
        # print("\n")
        # for i in range(len(file_name_list)):
        #     print(f"fnl[{i}]: ", file_name_list[i])
        # print("\n")

        # need_to_remove = list(set(file_name_list).difference(file_listing)) #returned result = what IS in sqlitedb table, and NOT IN dir-file-listing
        # need_to_add = list(set(file_listing).difference(file_name_list)) #returned result = what IS in dir-file-listing, and NOT IN sqlitedb table

        need_to_remove = []
        need_to_add = []

        # file_name_list = in database
        # file_listing = directory scan

        #add to need_to_remove list if item in database is not in directory
        for index, i in enumerate(file_name_list):
            if i not in [x[0] for x in file_listing]:
                need_to_remove.append(i)
                self.tableWidget_2.setItem(2, 0, QTableWidgetItem(f"Scanned {index + 1} out of {len(file_name_list)} file(s)"))
                self.progressBar_6.setValue((index + 1) // (len(file_name_list))*100)
            # elif i in file_listing:
            #     need_to_add.append(i)
            #     self.tableWidget_2.setItem(3, 0, QTableWidgetItem(f"Found {index + 1} new file(s) to add"))
            #     self.progressBar_5.setValue((index + 1) // (len(file_listing))*100)
        
        # print("\n")
        # for i in range(len(file_listing)):
        #     print(f"FL[{i}]: ", file_listing[i])
        # print("\n")

        #remove the 'need to remove' entries from the sqlitedb table
        for index, i in enumerate(need_to_remove):
            exe_str = f"DELETE FROM main_table_db where \"SOP name\" = \"{i}\""
            connection.execute(exe_str)
            self.tableWidget_2.setItem(2, 0, QTableWidgetItem(f"Removed {index + 1} out of {len(need_to_remove)} dead entries"))
            self.progressBar_6.setValue((index + 1) // (len(need_to_remove))*100)

        if len(need_to_remove) == 0:
            self.tableWidget_2.setItem(2, 0, QTableWidgetItem(f"No dead entries in table"))

        # add to need_to_add list if item in database is not in directory
        for index, i in enumerate(file_listing):
            if i[0] not in file_name_list:
                need_to_add.append(i)
                self.tableWidget_2.setItem(3, 0, QTableWidgetItem(f"Found {index + 1} new file(s) to add"))
                self.progressBar_5.setValue((index + 1) // (len(file_listing))*100)
        
        if len(need_to_add) == 0:
            self.tableWidget_2.setItem(3, 0, QTableWidgetItem(f"No new files in folder"))
            self.tableWidget_2.setItem(4, 0, QTableWidgetItem(f"Didn't need to add entries to database"))

        # print("\n")
        # for i in range(len(need_to_remove)):
        #     print(f"NTR[{i}]: ", need_to_remove[i])
        # print("\n")

        # print("\n")
        # for i in range(len(need_to_add)):
        #     print(f"NTA[{i}]: ", need_to_add[i])
        # print("\n")

        box_str = ""
        for index, i in enumerate(need_to_add):
            box_str += f"{index+1}.  {i[0]},\n -{i[1]},\n -{i[2]}\n"
        self.plainTextEdit.setPlainText(box_str)

        #add the 'need to add' entries to the sqlitedb table
        for index, i in enumerate(need_to_add):
            try:
                rev_num, issue_date, review_date, sop_code = self.get_SOP_info_from_doc(file=i[2])
            except:
                continue
            try:
                file_date = datetime.strptime(issue_date, '%d/%m/%Y')
                
                if datetime.now() > (file_date + relativedelta(months=+(mywindow.file_expiry))):
                    expiry_status = "expired"
                else:
                    expiry_status = "current"
            except:
                expiry_status = "'SOP issue date' info not found in document"

            exe_str = f"INSERT INTO main_table_db values(\"{sop_code}\",\"{i[0]}\",\"{i[1]}\",\"{rev_num}\",\"{issue_date}\",\"{review_date}\",\"{expiry_status}\",\"{i[2]}\",\"\",\"\")"
            connection.execute(exe_str)
            self.tableWidget_2.setItem(4, 0, QTableWidgetItem(f"Added {index + 1} out of {len(need_to_add)} file(s)"))
            self.progressBar.setValue((index + 1) // (len(need_to_add))*100)

        connection.commit()
        connection.close()

        mywindow.fill_table_sqlite(filter_params=None)
        self.label_2.setText("All done - Close this window by clicking the [ X ] button in the top-right corner")
        # self.close()


class MySplashScreen(QSplashScreen):
    def __init__(self, animation, flags):
        # run event dispatching in another thread
        QSplashScreen.__init__(self, QPixmap(), flags)
        self.movie = QMovie(animation)
        self.movie.frameChanged.connect(self.onNextFrame)
        # self.connect(self.movie, SIGNAL('frameChanged(int)'), SLOT('onNextFrame()'))
        self.movie.start()

    @pyqtSlot()
    def onNextFrame(self):
        pixmap = self.movie.currentPixmap()
        self.setPixmap(pixmap)
        self.setMask(pixmap.mask())


# ====== Functionality code ends here ======

if __name__ == '__main__':
    
    app = QApplication(sys.argv)

    splash = MySplashScreen('C:/Users/drift/OneDrive/PROJECTS/Mappe/icons/loading.gif', Qt.WindowStaysOnTopHint)
    splash.show()
    # QTimer.singleShot(2000, splash.close)

    mywindow = MainWindow()
    # mywindow = Window()
    
    mywindow.show()
    splash.close()
    
    sys.exit(app.exec_())
